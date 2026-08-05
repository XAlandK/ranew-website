"""
RanewWords.py
─────────────────────────
Interactive command-line tool that applies the taskpane's word-processing
actions to a .docx file.

Supported actions:
    1. Correct the Wrong Words
    2. Highlight the Skipped Words

This script is optimized for large documents and keeps the workflow simple:
    - Processes paragraphs, tables, headers, and footers
    - Replaces text inside runs to preserve formatting where possible
    - Uses interactive prompts for the same kinds of settings the taskpane exposes

Correction records (the "wrong"/"right" word pairs) are fetched from the
Supabase "CorrectedWords" table via supabase_repository.py, instead of a
local JSON file. See supabase_repository.py for the SUPABASE_URL /
SUPABASE_KEY environment variables used in production.

Requirements:
    pip install python-docx supabase

Usage:
    python "RanewWords.py"
"""

from __future__ import annotations

import json
import re
import sys
import time
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Iterable, Iterator, List, Optional, Sequence, Tuple

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor
    from docx.text.run import Run
except ImportError:
    print("Error: python-docx is not installed.")
    print("Install it with: pip install python-docx")
    sys.exit(1)

from supabase_repository import (
    CORRECTIONS_TABLE,
    CorrectionsRepository,
    SupabaseConnectionError,
    SupabaseQueryError,
)


SCRIPT_DIR = Path(__file__).resolve().parent
TASKPANE_DIR = SCRIPT_DIR.parent
DEFAULT_SKIPPED_WORDS = TASKPANE_DIR / "skipped_words.json"
DEFAULT_VOLUMES_DIR = SCRIPT_DIR / "Volumes"
DEFAULT_COMPLETED_DIR = SCRIPT_DIR / "CompletedVolumes"
DEFAULT_LOGS_DIR = SCRIPT_DIR / "VolumeLogs"
DEFAULT_OUTPUT_SUFFIX = "_completed"

# Broad Arabic diacritics/harakat coverage (tashkeel + Quranic marks).
ARABIC_HARAKAT_PATTERN = re.compile(r"[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED\u08D4-\u08FF]")
ARABIC_BASE_LETTER_PATTERN = re.compile(r"[\u0621-\u063A\u0641-\u064A\u066E\u066F\u0671-\u06D3\u06FA-\u06FF]")
ARABIC_WORD_CHAR_PATTERN = re.compile(
    r"[\u0621-\u063A\u0641-\u064A\u066E\u066F\u0671-\u06D3\u06FA-\u06FF\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED\u08D4-\u08FF]"
)
# Treat joiners and tatweel as in-word characters to avoid suffix false positives
# (for example when a word contains ZWNJ: U+200C).
WHOLE_WORD_BOUNDARY_CLASS = r"[\w\u0640\u200C\u200D]"
WHOLE_WORD_TOKEN_PATTERN = re.compile(rf"{WHOLE_WORD_BOUNDARY_CLASS}+")
# Any run of whitespace inside a multi-word rule should match any common
# whitespace / zero-width separator that Word documents place between words
# (regular space, tab, NBSP, narrow NBSP, ZWSP, ZWNJ, ZWJ, word joiner).
FLEXIBLE_WHITESPACE_PATTERN = r"[\s\u00A0\u200B\u200C\u200D\u202F\u2060]+"

# Narrow no-break space (U+202F) is normalized to a regular no-break space
# (U+00A0) everywhere, both in the loaded JSON rules and in the document text,
# so the two characters are never treated as different separators.
NARROW_NBSP = "\u202F"
NBSP = "\u00A0"

# Corrections are applied to each paragraph repeatedly until it stops changing,
# because one rule's output can become another rule's input (for example an
# Arabic-yeh normalization that then enables a multi-word join). This caps the
# number of sweeps so unexpected oscillating rules can never loop forever.
MAX_CORRECTION_PASSES = 5

# A "ـدا" token that comes right after a closing parenthesis ")" is an endorsed
# spelling in that context (the suffix attaches to the bracketed quotation) and
# must not be highlighted by the skipped-word pass. The capturing group isolates
# the "ـدا" span so only that token is protected, not the parenthesis. Optional
# zero-width / whitespace separators between ")" and "ـدا" are tolerated.
PAREN_SUFFIX_DA = "\u0640\u062f\u0627"  # ـدا (tatweel + dal + alef)
PAREN_SUFFIX_DA_PATTERN = re.compile(
    r"\)(?:" + FLEXIBLE_WHITESPACE_PATTERN + r")?(" + re.escape(PAREN_SUFFIX_DA) + r")"
)


def normalize_narrow_nbsp(text: str) -> str:
    if not text:
        return text
    return text.replace(NARROW_NBSP, NBSP)


WD_COLOR_MAP = {
    "YELLOW": WD_COLOR_INDEX.YELLOW,
    "BRIGHT_GREEN": WD_COLOR_INDEX.BRIGHT_GREEN,
    "TURQUOISE": WD_COLOR_INDEX.TURQUOISE,
    "PINK": WD_COLOR_INDEX.PINK,
    "BLUE": WD_COLOR_INDEX.BLUE,
    "RED": WD_COLOR_INDEX.RED,
    "DARK_BLUE": WD_COLOR_INDEX.DARK_BLUE,
    "TEAL": WD_COLOR_INDEX.TEAL,
    "GREEN": WD_COLOR_INDEX.GREEN,
    "VIOLET": WD_COLOR_INDEX.VIOLET,
    "DARK_RED": WD_COLOR_INDEX.DARK_RED,
    "DARK_YELLOW": WD_COLOR_INDEX.DARK_YELLOW,
    "GRAY_50": WD_COLOR_INDEX.GRAY_50,
    "GRAY_25": WD_COLOR_INDEX.GRAY_25,
    "BLACK": WD_COLOR_INDEX.BLACK,
    "WHITE": WD_COLOR_INDEX.WHITE,
    "AUTO": WD_COLOR_INDEX.AUTO,
}


class SessionLogger:
    """Mirror everything written to a stream into a session log file (tee)."""

    def __init__(self, stream, log_file):
        self._stream = stream
        self._log_file = log_file

    def write(self, data):
        self._stream.write(data)
        try:
            self._log_file.write(data)
        except Exception:
            pass
        return len(data)

    def flush(self):
        self._stream.flush()
        try:
            self._log_file.flush()
        except Exception:
            pass

    def isatty(self):
        return getattr(self._stream, "isatty", lambda: False)()


def start_session_log(logs_dir: Path):
    """Create VolumeLogs/<timestamp>.txt and tee stdout/stderr into it.

    Returns (log_file, original_stdout, original_stderr, log_path) so the caller
    can restore the streams and close the file when the session ends.
    """
    logs_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    log_path = logs_dir / f"session_{timestamp}.txt"
    log_file = open(log_path, "w", encoding="utf-8")
    log_file.write(f"Ranew Words session log - {datetime.now().isoformat(timespec='seconds')}\n")
    log_file.write("=" * 60 + "\n")
    log_file.flush()

    original_stdout = sys.stdout
    original_stderr = sys.stderr
    sys.stdout = SessionLogger(original_stdout, log_file)
    sys.stderr = SessionLogger(original_stderr, log_file)
    return log_file, original_stdout, original_stderr, log_path


def stop_session_log(log_file, original_stdout, original_stderr) -> None:
    sys.stdout = original_stdout
    sys.stderr = original_stderr
    try:
        log_file.flush()
        log_file.close()
    except Exception:
        pass


class FormattingOptions:
    def __init__(
        self,
        font_family: Optional[str] = None,
        font_size: Optional[float] = None,
        bold: bool = False,
        color_hex: Optional[str] = None,
    ):
        self.font_family = font_family
        self.font_size = font_size
        self.bold = bold
        self.color_hex = color_hex


def prompt_text(message: str, default: Optional[str] = None) -> str:
    suffix = f" [{default}]" if default is not None else ""
    value = input(f"{message}{suffix}: ").strip()
    if not value and default is not None:
        return default
    return value


def prompt_yes_no(message: str, default: bool = True) -> bool:
    default_hint = "Y/n" if default else "y/N"
    while True:
        value = input(f"{message} ({default_hint}): ").strip().lower()
        if not value:
            return default
        if value in {"y", "yes"}:
            return True
        if value in {"n", "no"}:
            return False
        print("Please answer yes or no.")


def prompt_choice(message: str, options: Sequence[Tuple[str, str]]) -> str:
    print(message)
    for key, label in options:
        print(f"  {key}. {label}")
    valid = {key for key, _ in options}
    while True:
        value = input("Select an option: ").strip()
        if value in valid:
            return value
        print("Please select one of the listed options.")


def prompt_highlight_choice(message: str, default: str) -> str:
    options = [
        ("YELLOW", "Yellow"),
        ("BRIGHT_GREEN", "Bright Green"),
        ("TURQUOISE", "Turquoise"),
        ("PINK", "Pink"),
        ("BLUE", "Blue"),
        ("RED", "Red"),
        ("DARK_BLUE", "Dark Blue"),
        ("TEAL", "Teal"),
        ("GREEN", "Green"),
        ("VIOLET", "Violet"),
        ("DARK_RED", "Dark Red"),
        ("DARK_YELLOW", "Dark Yellow"),
        ("GRAY_50", "Gray 50%"),
        ("GRAY_25", "Gray 25%"),
        ("BLACK", "Black"),
        ("AUTO", "Auto / clear"),
    ]
    print(message)
    for key, label in options:
        marker = " (default)" if key == default else ""
        print(f"  {key}. {label}{marker}")
    valid = {key for key, _ in options}
    while True:
        value = input(f"Select a highlight color [{default}]: ").strip().upper()
        if not value:
            return default
        if value in valid:
            return value
        print("Please select one of the listed highlight colors.")


def prompt_light_highlight_choice(message: str, default: str, disallow: Optional[str] = None) -> str:
    options = [
        ("YELLOW", "Yellow (light)"),
        ("BRIGHT_GREEN", "Bright Green (light)"),
        ("TURQUOISE", "Turquoise (light)"),
        ("PINK", "Pink (light)"),
        ("GRAY_25", "Gray 25% (light)"),
    ]

    if disallow:
        options = [opt for opt in options if opt[0] != disallow]

    print(message)
    for key, label in options:
        marker = " (default)" if key == default else ""
        print(f"  {key}. {label}{marker}")

    valid = {key for key, _ in options}
    fallback = default if default in valid else next(iter(valid))

    while True:
        value = input(f"Select a light highlight color [{fallback}]: ").strip().upper()
        if not value:
            return fallback
        if value in valid:
            return value
        print("Please select one of the listed light highlight colors.")


def prompt_formatting_options() -> FormattingOptions:
    print("Formatting options for inserted/replaced text (leave blank to keep the original run's formatting):")
    font_family = prompt_text("  Font family", "") or None
    font_size_raw = prompt_text("  Font size", "")
    font_size = None
    if font_size_raw:
        try:
            font_size = float(font_size_raw)
        except ValueError:
            print("  Invalid size; leaving font size unchanged.")

    bold = prompt_yes_no("  Bold text?", default=False)
    color_hex = prompt_text("  Font color as hex (for example #000000)", "") or None
    if color_hex and not color_hex.startswith("#"):
        color_hex = f"#{color_hex}"

    return FormattingOptions(
        font_family=font_family,
        font_size=font_size,
        bold=bold,
        color_hex=color_hex,
    )


def load_json_list(path: Path) -> list:
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"Warning: could not read JSON from {path}: {exc}")
        return []
    return data if isinstance(data, list) else []


def load_skipped_words_from_file(path: Path) -> list[str]:
    raw = load_json_list(path)
    # Ignore 1-character entries to avoid standalone-letter highlighting.
    words = [normalize_narrow_nbsp(str(word).strip()) for word in raw if len(str(word).strip()) >= 2]
    return sorted(set(words), key=len, reverse=True)


def has_skipped_words(path: Path) -> bool:
    return len(load_skipped_words_from_file(path)) > 0


def normalize_docx_path(value: str) -> Path:
    path = Path(value.strip().strip('"'))
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    if not path.is_absolute():
        path = (SCRIPT_DIR / path).resolve()
    return path


def normalize_output_path(input_path: Path, value: str) -> Path:
    path = Path(value.strip().strip('"'))
    # python-docx writes OOXML documents, so force .docx for reliability.
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    if not path.is_absolute():
        path = (input_path.parent / path).resolve()
    return path


def iter_paragraphs(container) -> Iterator:
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def iter_all_paragraphs(document) -> Iterator:
    yield from iter_paragraphs(document)
    for section in document.sections:
        yield from iter_paragraphs(section.header)
        yield from iter_paragraphs(section.footer)


def contains_any(text: str, needles: Sequence[str]) -> bool:
    return any(needle and needle in text for needle in needles)


def compile_search_pattern(term: str) -> re.Pattern:
    normalized = term.strip()
    if not normalized:
        return re.compile(r"a^")

    escaped = re.escape(normalized)
    if " " in normalized or any(ch in normalized for ch in "[](){}"):
        # Multi-token terms: let the space(s) in the rule match any run of common
        # whitespace / zero-width separators the document might use between words.
        # Without this a rule like "لە بارەی" never matches when the document
        # separates the words with NBSP, narrow NBSP, ZWNJ, a tab, etc.
        parts = re.split(r"\s+", normalized)
        return re.compile(FLEXIBLE_WHITESPACE_PATTERN.join(re.escape(part) for part in parts))

    return re.compile(rf"(?<!\w){escaped}(?!\w)")


def compile_whole_word_pattern(term: str) -> re.Pattern:
    normalized = term.strip()
    if not normalized:
        return re.compile(r"a^")

    escaped = re.escape(normalized)
    return re.compile(rf"(?<!{WHOLE_WORD_BOUNDARY_CLASS}){escaped}(?!{WHOLE_WORD_BOUNDARY_CLASS})")


def apply_run_formatting(run, formatting: FormattingOptions) -> None:
    if formatting.font_family:
        run.font.name = formatting.font_family
    if formatting.font_size:
        run.font.size = Pt(formatting.font_size)
    run.bold = formatting.bold
    if formatting.color_hex:
        hex_value = formatting.color_hex.lstrip("#")
        if len(hex_value) == 6:
          run.font.color.rgb = RGBColor.from_string(hex_value.upper())


def apply_highlight(run, highlight_key: str, skip_if_harakat: bool = True) -> None:
    if skip_if_harakat and contains_arabic_harakat(run.text):
        return

    if highlight_key == "AUTO":
        run.font.highlight_color = None
        return
    run.font.highlight_color = WD_COLOR_MAP[highlight_key]


def contains_arabic_harakat(text: str) -> bool:
    return bool(ARABIC_HARAKAT_PATTERN.search(text or ""))


def clear_highlight_by_color(document, highlight_key: str) -> int:
    if highlight_key == "AUTO":
        return 0

    target_color = WD_COLOR_MAP.get(highlight_key)
    if target_color is None:
        return 0

    cleared = 0
    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            if run.font.highlight_color == target_color:
                run.font.highlight_color = None
                cleared += 1
    return cleared


def remove_harakat_highlights(document) -> int:
    removed = 0

    for paragraph in iter_all_paragraphs(document):
        runs = list(paragraph.runs)
        if not runs:
            continue

        run_spans: list[Tuple[int, int, Run]] = []
        full_text_parts: list[str] = []
        cursor = 0

        for run in runs:
            text = run.text or ""
            if not text:
                continue
            start = cursor
            end = start + len(text)
            run_spans.append((start, end, run))
            full_text_parts.append(text)
            cursor = end

        if not run_spans:
            continue

        full_text = "".join(full_text_parts)
        n = len(full_text)
        i = 0

        while i < n:
            if not ARABIC_WORD_CHAR_PATTERN.match(full_text[i]):
                i += 1
                continue

            start = i
            has_harakat = False
            has_base_letter = False

            while i < n and ARABIC_WORD_CHAR_PATTERN.match(full_text[i]):
                ch = full_text[i]
                if ARABIC_HARAKAT_PATTERN.match(ch):
                    has_harakat = True
                if ARABIC_BASE_LETTER_PATTERN.match(ch):
                    has_base_letter = True
                i += 1

            end = i
            if not (has_harakat and has_base_letter):
                continue

            for run_start, run_end, run in run_spans:
                if run_end <= start or run_start >= end:
                    continue
                if run.font.highlight_color is not None:
                    run.font.highlight_color = None
                    removed += 1

    return removed


def replace_in_run_text(run, replacements: Sequence[Tuple[re.Pattern, str]], formatting: Optional[FormattingOptions] = None, highlight_key: Optional[str] = None) -> int:
    text = run.text
    if not text:
        return 0

    changed = 0
    for pattern, replacement in replacements:
        text, count = pattern.subn(replacement, text)
        changed += count

    if changed:
        run.text = text
        if formatting:
            apply_run_formatting(run, formatting)
        if highlight_key:
            apply_highlight(run, highlight_key)

    return changed


def process_paragraphs(document, replacements: Sequence[Tuple[re.Pattern, str]], formatting: Optional[FormattingOptions] = None, highlight_key: Optional[str] = None) -> int:
    total = 0
    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            total += replace_in_run_text(run, replacements, formatting=formatting, highlight_key=highlight_key)
    return total


def _is_whitespace_like(ch: Optional[str]) -> bool:
    if not ch:
        return True
    return ch.isspace() or ch == "\u00A0"


def _prev_non_empty_char(runs, index: int) -> Optional[str]:
    j = index - 1
    while j >= 0:
        t = runs[j].text or ""
        if t:
            return t[-1]
        j -= 1
    return None


def _next_non_empty_run(runs, index: int):
    j = index + 1
    while j < len(runs):
        t = runs[j].text or ""
        if t:
            return runs[j]
        j += 1
    return None


def separate_waw(document) -> int:
    total = 0
    waw = "و"
    commas = ("،", ",")

    # Strict rule requested: "و،" (or "و,") should become "و".
    # Work at the paragraph level so the comma is removed even when the
    # waw and the comma live in different runs, empty runs sit between them,
    # or the waw is part of a longer run (for example "...شتو،").
    for paragraph in iter_all_paragraphs(document):
        full_text, run_spans = get_paragraph_text_and_run_spans(paragraph)
        if not full_text or not run_spans:
            continue

        # Collect global indices of comma characters that follow a waw.
        delete_positions: list[int] = []
        for i in range(1, len(full_text)):
            if full_text[i] in commas and full_text[i - 1] == waw:
                delete_positions.append(i)

        if not delete_positions:
            continue

        # Map each global comma position to (run, local index).
        per_run_locals: dict[int, list[int]] = {}
        run_lookup: dict[int, Run] = {}
        for span_idx, (run_start, run_end, run) in enumerate(run_spans):
            run_lookup[span_idx] = run
            for pos in delete_positions:
                if run_start <= pos < run_end:
                    per_run_locals.setdefault(span_idx, []).append(pos - run_start)

        # Delete commas per run, in reverse order, to keep indices valid.
        for span_idx, local_indices in per_run_locals.items():
            run = run_lookup[span_idx]
            text = run.text or ""
            for local_index in sorted(set(local_indices), reverse=True):
                if 0 <= local_index < len(text) and text[local_index] in commas:
                    text = text[:local_index] + text[local_index + 1 :]
                    total += 1
            run.text = text

    return total


def normalize_document_narrow_nbsp(document) -> int:
    """Replace every narrow no-break space (U+202F) with NBSP (U+00A0) in run text."""
    total = 0
    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            text = run.text or ""
            if NARROW_NBSP in text:
                count = text.count(NARROW_NBSP)
                run.text = text.replace(NARROW_NBSP, NBSP)
                total += count
    return total


def load_dictionary_from_supabase(repository: CorrectionsRepository) -> list[dict[str, str]]:
    """Fetch correction records via `repository` and normalize them into the
    same [{"wrong": ..., "right": ...}, ...] shape the rest of the script
    (build_replacement_rules, apply_corrections, etc.) already expects."""
    try:
        raw_entries = repository.get_dictionary_entries()
    except (SupabaseConnectionError, SupabaseQueryError) as exc:
        print(f"Error: could not load corrections from Supabase: {exc}", file=sys.stderr)
        return []

    entries: list[dict[str, str]] = []
    for entry in raw_entries:
        wrong = normalize_narrow_nbsp(str(entry.get("wrong") or "").strip())
        right = normalize_narrow_nbsp(str(entry.get("right") or "").strip())
        if wrong and right:
            entries.append({"wrong": wrong, "right": right})
    return entries


def build_replacement_rules(entries: Sequence[dict[str, str]]) -> list[Tuple[re.Pattern, str, str]]:
    ordered = sorted(entries, key=lambda item: len(item["wrong"]), reverse=True)
    rules: list[Tuple[re.Pattern, str, str]] = []
    seen = set()
    for entry in ordered:
        wrong = entry["wrong"].strip()
        right = entry["right"]
        if wrong in seen:
            continue
        seen.add(wrong)
        rules.append((compile_search_pattern(wrong), wrong, right))
    return rules


def apply_corrections(document, rules: Sequence[Tuple[re.Pattern, str, str]], formatting: Optional[FormattingOptions], highlight_key: Optional[str]) -> int:
    paragraphs = list(iter_all_paragraphs(document))
    total_paragraphs = len(paragraphs)

    rule_index_one: dict[str, list[Tuple[re.Pattern, str, str]]] = {}
    rule_index_two: dict[str, list[Tuple[re.Pattern, str, str]]] = {}
    for pattern, wrong, right in rules:
        if not wrong:
            continue
        # Bucket by the first whitespace-delimited token so multi-word rules
        # (e.g. "لە بارەی") are keyed on their leading letters, independent of
        # whatever separator the document uses between the words.
        first_token = wrong.split()[0]
        if len(first_token) >= 2:
            rule_index_two.setdefault(first_token[:2], []).append((pattern, wrong, right))
        else:
            rule_index_one.setdefault(first_token[0], []).append((pattern, wrong, right))

    print(f"  Starting correction pass: {total_paragraphs} paragraphs, {len(rules)} rules.")
    print("  Progress updates will appear every 25 paragraphs.")

    total = 0
    start_time = time.time()
    for idx, paragraph in enumerate(paragraphs, start=1):
        # Operate on the full paragraph text (not run-by-run) so that words split
        # across multiple runs (e.g. "بەرز" + "و" forming "بەرزو") are still matched.
        paragraph_text, _ = get_paragraph_text_and_run_spans(paragraph)
        if not paragraph_text:
            continue

        candidate_rules: list[Tuple[re.Pattern, str, str]] = []

        if len(paragraph_text) >= 2:
            two_keys = {paragraph_text[i : i + 2] for i in range(len(paragraph_text) - 1)}
            for key in two_keys:
                bucket = rule_index_two.get(key)
                if bucket:
                    candidate_rules.extend(bucket)

        one_keys = set(paragraph_text)
        for key in one_keys:
            bucket = rule_index_one.get(key)
            if bucket:
                candidate_rules.extend(bucket)

        if not candidate_rules:
            continue

        # Apply longest wrong-forms first so that shorter rules cannot pre-empt
        # a longer, more specific correction.
        candidate_rules.sort(key=lambda item: len(item[1]), reverse=True)

        total += apply_paragraph_corrections(paragraph, candidate_rules, formatting, highlight_key)

        if idx % 25 == 0 or idx == total_paragraphs:
            elapsed = int(time.time() - start_time)
            pct = int((idx / total_paragraphs) * 100) if total_paragraphs else 100
            print(f"  Progress: {idx}/{total_paragraphs} paragraphs ({pct}%) | replacements: {total} | {elapsed}s")

    elapsed = int(time.time() - start_time)
    print(f"  Correction phase finished in {elapsed}s.")
    return total


def replace_span_in_runs(runs: Sequence[Run], start: int, end: int, replacement: str, changed: set) -> None:
    """Replace the global character span [start, end) of the concatenated run text.

    The replacement is written into the first run that overlaps the span (keeping
    its physical position and formatting); any text the span covers in following
    runs is removed. Run start/end offsets are computed from the run lengths as
    they exist at call time, so callers must process the matches of a single rule
    from right to left to keep earlier (left-side) coordinates valid.
    """
    cursor = 0
    first_run = None
    for run in runs:
        run_text = run.text or ""
        run_len = len(run_text)
        run_start = cursor
        run_end = cursor + run_len
        cursor = run_end

        if run_end <= start or run_start >= end:
            continue

        local_start = max(start, run_start) - run_start
        local_end = min(end, run_end) - run_start

        if first_run is None:
            run.text = run_text[:local_start] + replacement + run_text[local_end:]
            first_run = run
        else:
            run.text = run_text[:local_start] + run_text[local_end:]
        changed.add(run)


def apply_paragraph_corrections(
    paragraph,
    candidate_rules: Sequence[Tuple[re.Pattern, str, str]],
    formatting: Optional[FormattingOptions],
    highlight_key: Optional[str],
) -> int:
    """Apply correction rules across the whole paragraph, spanning runs when needed."""
    runs = [run for run in paragraph.runs if (run.text or "")]
    if not runs:
        return 0

    changed: set = set()
    replacements = 0
    text = "".join(run.text or "" for run in runs)

    # Apply the candidate rules repeatedly until the paragraph stops changing.
    # A single ordered pass is not enough because one rule's output can become
    # another rule's input. For example, when the document uses an Arabic yeh
    # (ى U+0649) the normalization rule "بارەى" -> "بارەی" must run before the
    # multi-word rule "لە بارەی" -> "لەبارەی" can match. Since rules are applied
    # longest-first, the multi-word rule is tried first and misses the Arabic
    # spelling, so we loop until convergence (bounded to avoid any oscillation).
    for _ in range(MAX_CORRECTION_PASSES):
        pass_replacements = 0
        for pattern, wrong, right in candidate_rules:
            # Single-token rules can use a cheap substring check to skip the regex.
            # Multi-word rules must not, because their whitespace can be any separator
            # in the document, so the literal "wrong" may never appear verbatim.
            if " " not in wrong and wrong not in text:
                continue
            matches = list(pattern.finditer(text))
            if not matches:
                continue
            for match in reversed(matches):
                replace_span_in_runs(runs, match.start(), match.end(), right, changed)
                pass_replacements += 1
            text = "".join(run.text or "" for run in runs)
        replacements += pass_replacements
        if pass_replacements == 0:
            break

    if changed and (formatting or highlight_key):
        for run in changed:
            if formatting:
                apply_run_formatting(run, formatting)
            if highlight_key:
                apply_highlight(run, highlight_key)

    return replacements


def collect_right_word_patterns(entries: Sequence[dict[str, str]]) -> list[re.Pattern]:
    patterns: list[re.Pattern] = []
    seen = set()
    for entry in entries:
        right = str(entry.get("right") or "").strip()
        if not right or right in seen:
            continue
        seen.add(right)
        patterns.append(compile_search_pattern(right))
    patterns.sort(key=lambda pattern: len(pattern.pattern), reverse=True)
    return patterns


def get_paragraph_text_and_run_spans(paragraph) -> Tuple[str, list[Tuple[int, int, Run]]]:
    run_spans: list[Tuple[int, int, Run]] = []
    parts: list[str] = []
    cursor = 0

    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        start = cursor
        end = start + len(text)
        run_spans.append((start, end, run))
        parts.append(text)
        cursor = end

    return "".join(parts), run_spans


def compile_phrase_with_token_groups(phrase: str) -> re.Pattern:
    """Like compile_search_pattern for a multi-word phrase, but wraps every token
    in its own capturing group so the caller can recover each token's character
    span. Whitespace between tokens still matches any common separator."""
    parts = re.split(r"\s+", phrase.strip())
    grouped = FLEXIBLE_WHITESPACE_PATTERN.join("(" + re.escape(part) + ")" for part in parts)
    return re.compile(grouped)


def build_protected_skip_phrases(
    dictionary_entries: Sequence[dict[str, str]],
    skipped_words: Sequence[str],
) -> list[Tuple[re.Pattern, list[int], str]]:
    """Build the multi-word corrected (right-hand) phrases whose non-leading
    token(s) are themselves skipped words.

    These right-hand forms are spellings the correction step deliberately
    produces (for example "مەبەست پێی"). A skipped word that appears as a
    trailing token of such a phrase — i.e. with the rest of the phrase in front
    of it — is therefore already endorsed in that context and must not be
    re-flagged by the skipped-word highlighter.

    The rule is fully data-driven: it derives only from the correction dictionary
    and the skipped-word list, so it adapts to any phrase without hard-coding
    specific words. Each result is (token-group pattern, protected group numbers,
    first token) where a protected group number g matches the g-th token's span.
    """
    skipped_set = set(skipped_words)
    phrases: list[Tuple[re.Pattern, list[int], str]] = []
    seen: set[str] = set()
    for entry in dictionary_entries:
        right = normalize_narrow_nbsp(str(entry.get("right") or "").strip())
        if not right or " " not in right or right in seen:
            continue
        tokens = re.split(r"\s+", right)
        if len(tokens) < 2:
            continue
        # Only protect non-leading tokens, so a skipped word is suppressed only
        # when it follows its phrase context (e.g. "پێی" after "مەبەست"), never
        # when it appears on its own elsewhere.
        protected_groups = [i + 1 for i in range(1, len(tokens)) if tokens[i] in skipped_set]
        if not protected_groups:
            continue
        seen.add(right)
        phrases.append((compile_phrase_with_token_groups(right), protected_groups, tokens[0]))
    return phrases


def highlight_words(
    document,
    highlight_key: str,
    skipped_words_path: Path = DEFAULT_SKIPPED_WORDS,
    dictionary_entries: Sequence[dict[str, str]] = (),
) -> int:
    unique_words = load_skipped_words_from_file(skipped_words_path)
    if not unique_words:
        print(f"  No skipped words found in {skipped_words_path}")
        return 0

    single_words = [word for word in unique_words if " " not in word]
    multiword_terms = [word for word in unique_words if " " in word]

    pattern_pairs: list[Tuple[re.Pattern, str]] = [(compile_whole_word_pattern(word), word) for word in multiword_terms]
    single_word_set = set(single_words)

    word_index_one: dict[str, list[Tuple[re.Pattern, str]]] = {}
    word_index_two: dict[str, list[Tuple[re.Pattern, str]]] = {}
    for pattern, word in pattern_pairs:
        if len(word) >= 2:
            word_index_two.setdefault(word[:2], []).append((pattern, word))
        else:
            word_index_one.setdefault(word[0], []).append((pattern, word))

    # Multi-word corrected phrases whose trailing token is itself a skipped word
    # (e.g. "مەبەست پێی"). When such a phrase is present, its protected token must
    # not be highlighted even though it also appears in the skipped-word list.
    protected_phrases = build_protected_skip_phrases(dictionary_entries, unique_words)
    phrase_index_one: dict[str, list[Tuple[re.Pattern, list[int]]]] = {}
    phrase_index_two: dict[str, list[Tuple[re.Pattern, list[int]]]] = {}
    for pattern, groups, first_token in protected_phrases:
        if len(first_token) >= 2:
            phrase_index_two.setdefault(first_token[:2], []).append((pattern, groups))
        else:
            phrase_index_one.setdefault(first_token[0], []).append((pattern, groups))

    paragraphs = list(iter_all_paragraphs(document))
    total_paragraphs = len(paragraphs)

    print(f"  Starting skipped-word highlight pass: {total_paragraphs} paragraphs, {len(unique_words)} words from {skipped_words_path}.")
    print("  Progress is shown on one updating line.")

    total = 0
    start_time = time.time()
    for idx, paragraph in enumerate(paragraphs, start=1):
        paragraph_text, run_spans = get_paragraph_text_and_run_spans(paragraph)
        if not paragraph_text or not run_spans:
            continue

        candidate_pairs: list[Tuple[re.Pattern, str]] = []
        if len(paragraph_text) >= 2:
            two_keys = {paragraph_text[i : i + 2] for i in range(len(paragraph_text) - 1)}
            for key in two_keys:
                bucket = word_index_two.get(key)
                if bucket:
                    candidate_pairs.extend(bucket)

        one_keys = set(paragraph_text)
        for key in one_keys:
            bucket = word_index_one.get(key)
            if bucket:
                candidate_pairs.extend(bucket)

        global_spans: list[Tuple[int, int]] = []

        # Strict token matching for single words prevents partial/suffix highlights.
        for token_match in WHOLE_WORD_TOKEN_PATTERN.finditer(paragraph_text):
            token = token_match.group(0)
            if token in single_word_set:
                global_spans.append((token_match.start(), token_match.end()))

        for pattern, word in candidate_pairs:
            if word not in paragraph_text:
                continue
            for match in pattern.finditer(paragraph_text):
                global_spans.append((match.start(), match.end()))

        # Suppress any skipped-word span that falls inside an endorsed corrected
        # phrase (e.g. "پێی" within "مەبەست پێی"), so context-correct spellings
        # are not re-flagged by the highlighter.
        if global_spans:
            protected_spans: list[Tuple[int, int]] = []

            # "ـدا" right after a closing parenthesis ")" is endorsed in context.
            if PAREN_SUFFIX_DA in paragraph_text and ")" in paragraph_text:
                for match in PAREN_SUFFIX_DA_PATTERN.finditer(paragraph_text):
                    gstart, gend = match.start(1), match.end(1)
                    if gstart >= 0 and gend > gstart:
                        protected_spans.append((gstart, gend))

            protected_phrase_candidates: list[Tuple[re.Pattern, list[int]]] = []
            if len(paragraph_text) >= 2:
                for key in two_keys:
                    bucket = phrase_index_two.get(key)
                    if bucket:
                        protected_phrase_candidates.extend(bucket)
            for key in one_keys:
                bucket = phrase_index_one.get(key)
                if bucket:
                    protected_phrase_candidates.extend(bucket)

            if protected_phrase_candidates:
                for pattern, groups in protected_phrase_candidates:
                    for match in pattern.finditer(paragraph_text):
                        for group in groups:
                            gstart, gend = match.start(group), match.end(group)
                            if gstart >= 0 and gend > gstart:
                                protected_spans.append((gstart, gend))

            if protected_spans:
                global_spans = [
                    (span_start, span_end)
                    for span_start, span_end in global_spans
                    if not any(
                        protected_start <= span_start and span_end <= protected_end
                        for protected_start, protected_end in protected_spans
                    )
                ]

        if not global_spans:
            continue

        # Convert paragraph-level match spans to run-local spans.
        for run_start, run_end, run in run_spans:
            local_spans: list[Tuple[int, int]] = []
            for span_start, span_end in global_spans:
                if span_end <= run_start or span_start >= run_end:
                    continue
                local_start = max(span_start, run_start) - run_start
                local_end = min(span_end, run_end) - run_start
                if local_start < local_end:
                    local_spans.append((local_start, local_end))

            if not local_spans:
                continue

            highlighted_segments = highlight_run_spans(run, local_spans, highlight_key)
            total += highlighted_segments

        if idx % 100 == 0 or idx == total_paragraphs:
            elapsed = int(time.time() - start_time)
            pct = int((idx / total_paragraphs) * 100) if total_paragraphs else 100
            print(
                f"  Progress: {idx}/{total_paragraphs} paragraphs ({pct}%) | highlights: {total} | {elapsed}s",
                end="\r",
                flush=True,
            )

    elapsed = int(time.time() - start_time)
    print()
    print(f"  Skipped-word highlight phase finished in {elapsed}s.")
    return total


def highlight_run_spans(run, spans: Sequence[Tuple[int, int]], highlight_key: str) -> int:
    text = run.text
    if not text:
        return 0

    ordered = sorted(set(spans), key=lambda item: (item[0], item[1]))
    merged: list[Tuple[int, int]] = []
    for start, end in ordered:
        if start >= end:
            continue
        if not merged or start >= merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    if not merged:
        return 0

    segments: list[Tuple[str, bool]] = []
    cursor = 0
    for start, end in merged:
        if cursor < start:
            segments.append((text[cursor:start], False))
        segments.append((text[start:end], True))
        cursor = end
    if cursor < len(text):
        segments.append((text[cursor:], False))

    base_r = run._r
    parent_obj = run._parent
    # Snapshot the run's original XML BEFORE any highlight is applied. The first
    # segment reuses base_r (and may be mutated by apply_highlight), so every
    # subsequent segment must be cloned from this clean template instead of from
    # base_r; otherwise a highlight applied to an earlier segment bleeds into the
    # following ones (for example highlighting "باسم" would also color "کردووە").
    clean_template = deepcopy(base_r)
    prev_r = base_r
    base_used = False
    highlighted_count = 0

    for segment_text, should_highlight in segments:
        if not segment_text:
            continue

        if not base_used:
            target_r = base_r
            base_used = True
        else:
            target_r = deepcopy(clean_template)
            prev_r.addnext(target_r)

        segment_run = Run(target_r, parent_obj)
        segment_run.text = segment_text

        if should_highlight:
            # Match taskpane skipped-word behavior: no harakat-based exclusion.
            apply_highlight(segment_run, highlight_key, skip_if_harakat=False)
            highlighted_count += 1

        prev_r = target_r

    return highlighted_count


def clear_word_highlights(document, words: Sequence[str]) -> int:
    if not words:
        return 0

    patterns = [compile_search_pattern(word) for word in sorted(set(words), key=len, reverse=True)]
    paragraphs = list(iter_all_paragraphs(document))
    total_paragraphs = len(paragraphs)

    print(f"  Starting corrected-highlight clear pass: {total_paragraphs} paragraphs, {len(words)} words.")
    print("  Progress updates will appear every 100 paragraphs.")

    total = 0
    start_time = time.time()
    for idx, paragraph in enumerate(paragraphs, start=1):
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            if not contains_any(text, words):
                continue

            if any(pattern.search(text) for pattern in patterns):
                run.font.highlight_color = None
                total += 1

        if idx % 100 == 0 or idx == total_paragraphs:
            elapsed = int(time.time() - start_time)
            pct = int((idx / total_paragraphs) * 100) if total_paragraphs else 100
            print(f"  Progress: {idx}/{total_paragraphs} paragraphs ({pct}%) | cleared: {total} | {elapsed}s")

    elapsed = int(time.time() - start_time)
    print(f"  Corrected-highlight clear phase finished in {elapsed}s.")
    return total


def insert_wrong_words_list(document, wrong_words: Sequence[str], formatting: Optional[FormattingOptions] = None) -> None:
    if not wrong_words:
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run(" ".join(wrong_words))
    if formatting:
        apply_run_formatting(run, formatting)


def get_unique_wrong_words(entries: Sequence[dict[str, str]]) -> list[str]:
    seen = set()
    result = []
    for entry in entries:
        wrong = str(entry.get("wrong") or "").strip()
        if wrong and wrong not in seen:
            seen.add(wrong)
            result.append(wrong)
    return result


def get_unique_right_words(entries: Sequence[dict[str, str]]) -> list[str]:
    seen = set()
    result = []
    for entry in entries:
        right = str(entry.get("right") or "").strip()
        if right and right not in seen:
            seen.add(right)
            result.append(right)
    return result


def load_docx(path: Path):
    if not path.exists():
        raise FileNotFoundError(f"Input file not found: {path}")
    return Document(str(path))


def volume_sort_key(path: Path) -> Tuple[int, int, str]:
    # Sort numerically by the leading number in the filename (1, 2, ... 25, 30).
    # Files without a leading number are pushed to the end and sorted by name.
    match = re.match(r"\d+", path.stem)
    if match:
        return (0, int(match.group(0)), path.stem.lower())
    return (1, 0, path.stem.lower())


def find_volume_files(volumes_dir: Path) -> list[Path]:
    if not volumes_dir.exists():
        return []
    files = [
        path
        for path in volumes_dir.glob("*.docx")
        if not path.name.startswith("~$") and not path.stem.endswith(DEFAULT_OUTPUT_SUFFIX)
    ]
    return sorted(files, key=volume_sort_key)


def process_volume_document(
    input_path: Path,
    output_path: Path,
    dictionary_entries: Sequence[dict[str, str]],
    skipped_words_path: Path,
    correction_highlight_key: Optional[str],
    skipped_highlight_key: str,
) -> list[str]:
    """Apply both actions (corrections then skipped-word highlighting) to one document."""
    document = load_docx(input_path)
    total_changes: list[str] = []

    print("  [Pre-process] Normalizing narrow NBSP (U+202F) to NBSP (U+00A0)...")
    normalized_nbsp = normalize_document_narrow_nbsp(document)
    total_changes.append(f"Narrow NBSP normalized: {normalized_nbsp}")
    print(f"    Normalized {normalized_nbsp} narrow NBSP character(s).")

    print("  [WAW] Integrated prepass before corrections...")
    waw_count = separate_waw(document)
    total_changes.append(f"Waw fixes: {waw_count}")
    print(f"    Applied {waw_count} change(s).")

    print("  [1] Correcting wrong words...")
    rules = build_replacement_rules(dictionary_entries)
    corrections = apply_corrections(document, rules, None, correction_highlight_key)
    total_changes.append(f"Corrections: {corrections}")
    print(f"    Applied {corrections} run-level replacement(s).")

    print("  [2] Highlighting skipped words...")
    if has_skipped_words(skipped_words_path):
        skipped = highlight_words(document, skipped_highlight_key, skipped_words_path, dictionary_entries)
        total_changes.append(f"Skipped highlights: {skipped}")
        print(f"    Applied highlight to {skipped} run(s).")
    else:
        print(f"    No skipped words found in {skipped_words_path}")

    print("  [Post-process] Removing highlights from text containing Arabic diacritics...")
    removed_harakat = remove_harakat_highlights(document)
    total_changes.append(f"Harakat highlight cleanup: {removed_harakat}")
    print(f"    Removed highlights from {removed_harakat} run(s).")

    print("  [Post-process] Final Waw separation pass...")
    final_waw_count = separate_waw(document)
    total_changes.append(f"Final Waw fixes: {final_waw_count}")
    print(f"    Applied {final_waw_count} final Waw change(s).")

    document.save(str(output_path))
    return total_changes


def main() -> None:
    log_file, original_stdout, original_stderr, log_path = start_session_log(DEFAULT_LOGS_DIR)
    try:
        run_session(log_path)
    finally:
        stop_session_log(log_file, original_stdout, original_stderr)
        print(f"Session log saved to: {log_path}")


def run_session(log_path: Path) -> None:
    print("Ranew Words - Word file automation")
    print("This script edits every .docx file in the Volumes folder using the taskpane rules.")
    print(f"Session log: {log_path}")

    volume_files = find_volume_files(DEFAULT_VOLUMES_DIR)
    if not volume_files:
        print(f"Error: no .docx files found in {DEFAULT_VOLUMES_DIR}", file=sys.stderr)
        sys.exit(1)

    skipped_words_path = DEFAULT_SKIPPED_WORDS

    corrections_repository = CorrectionsRepository()
    dictionary_entries = load_dictionary_from_supabase(corrections_repository)
    skipped_word_count = len(load_skipped_words_from_file(skipped_words_path))

    print(f"Loaded corrected entries: {len(dictionary_entries)} from Supabase table '{CORRECTIONS_TABLE}'")
    print(f"Loaded skipped words: {skipped_word_count} from {skipped_words_path}")

    if not dictionary_entries:
        print("Error: no correction entries could be loaded from Supabase, so corrections cannot run.")
        sys.exit(1)

    print("\nVolumes to process (ascending order):")
    for path in volume_files:
        print(f"  - {path.name}")

    # Corrected words are never highlighted; only skipped words are.
    correction_highlight_key: Optional[str] = None
    skipped_highlight_key = prompt_highlight_choice(
        "Choose a highlight color for skipped words:",
        default="YELLOW",
    )

    DEFAULT_COMPLETED_DIR.mkdir(parents=True, exist_ok=True)

    for index, input_path in enumerate(volume_files, start=1):
        output_path = DEFAULT_COMPLETED_DIR / (
            f"{input_path.stem}{DEFAULT_OUTPUT_SUFFIX}{input_path.suffix}"
        )
        print(f"\n=== [{index}/{len(volume_files)}] Processing {input_path.name} ===")
        changes = process_volume_document(
            input_path,
            output_path,
            dictionary_entries,
            skipped_words_path,
            correction_highlight_key,
            skipped_highlight_key,
        )
        print(f"  Saved output to: {output_path}")
        for line in changes:
            print(f"    - {line}")

    print("\nDone. All volumes processed.")


if __name__ == "__main__":
    main()